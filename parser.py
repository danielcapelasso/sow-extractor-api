from docx import Document
from io import BytesIO
import re

def extract_sow_data(file_bytes):
    doc = Document(BytesIO(file_bytes))

    result = {
        "consultor_responsavel": None,
        "contexto_projeto": "",
        "etapas_projeto": None,
        "tempo_estimado": None,
        "principais_regras_negocio": [],
        "resumo_servicos": [],
        "casos_de_uso_detalhados": [],
        "casos_custom_resumo": [],
        "casos_custom_detalhados": []
    }

    current_section = None
    found_resumo_servicos = False
    found_custom_section = False

    # Folha resumo
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if len(cells) >= 2:
                raw_key = cells[0]
                key = re.sub(r"\s+", " ", raw_key).strip().lower()
                value = cells[1].strip()

                if not result["consultor_responsavel"] and ("preparado por" in key or "preparado pela" in key):
                    result["consultor_responsavel"] = value.split("/")[0].strip()
                elif not result["etapas_projeto"] and "etapas do projeto" in key:
                    result["etapas_projeto"] = value
                elif not result["tempo_estimado"] and "tempo estimado" in key:
                    result["tempo_estimado"] = value

    # Contexto e demais seções
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        if current_section != "contexto_projeto" and (
            "objetivo geral" in text.lower() or "objetivos de negócio mapeado" in text.lower()
        ):
            current_section = "contexto_projeto"
            continue

        if current_section != "principais_regras_negocio" and (
            "principais regras de negócio" in text.lower()
        ):
            current_section = "principais_regras_negocio"
            continue

        if not found_resumo_servicos and "resumo de serviços a serem providos" in text.lower():
            found_resumo_servicos = True
            current_section = None
            continue

        if any(
            t in text.lower()
            for t in [
                "abordagem", "gaps", "arquitetura", "esforços", "glossário", "anexo"
            ]
        ):
            current_section = None
            continue

        if current_section == "contexto_projeto":
            result["contexto_projeto"] += " " + text

    if result["contexto_projeto"]:
        result["contexto_projeto"] = result["contexto_projeto"].strip()
        result["contexto_projeto"] = re.sub(
            r"^(1\.\d\s*)?Objetivos? de negócio Mapeado\s*", "", result["contexto_projeto"], flags=re.IGNORECASE
        )

    # Resumo de Serviços
    if found_resumo_servicos:
        for table in doc.tables:
            header_cells = [cell.text.strip().lower() for cell in table.rows[0].cells]
            if "serviço" in header_cells or "servico" in header_cells:
                for row in table.rows[1:]:
                    cols = [cell.text.strip() for cell in row.cells]
                    servico = {}
                    for idx, col in enumerate(cols):
                        key = header_cells[idx] if idx < len(header_cells) else f"coluna_{idx+1}"
                        servico[key] = col
                    if any(servico.values()):
                        result["resumo_servicos"].append(servico)
                break

# Principais regras de negócio
    for table in doc.tables[1:]:
        headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
        for row in table.rows[1:]:
            cols = [cell.text.strip() for cell in row.cells]
            if len(cols) >= 5:
                if "descrição" in headers[2]:  # modelo novo
                    regra = {
                        "caso_uso": cols[0],
                        "regra_negocio": cols[1],
                        "descricao_regra": cols[2],
                        "comentario": cols[3],
                        "atendido": cols[4]
                    }
                else:  # modelo antigo
                    regra = {
                        "item": cols[0],
                        "regra_negocio": cols[1],
                        "atendido": cols[2],
                        "comentario": cols[3],
                        "caso_uso": cols[4]
                    }

                if any(v for v in regra.values()):
                    result["principais_regras_negocio"].append(regra)

    # Casos custom - RESUMO
    for table in doc.tables:
        header = [cell.text.strip().lower() for cell in table.rows[0].cells]
        if all(h in header for h in ["caso de uso", "tipo", "integração", "descrição do caso de uso"]):
            for row in table.rows[1:]:
                cols = [cell.text.strip() for cell in row.cells]
                if len(cols) >= 4:
                    result["casos_custom_resumo"].append({
                        "caso_de_uso": cols[0],
                        "tipo": cols[1],
                        "integracao": cols[2],
                        "descricao": cols[3]
                    })
   
# Casos custom - DETALHES
current_custom = None
section = None
collecting = False

for para in doc.paragraphs:
    text = para.text.strip()
    if not text:
        continue

    # Início da seção
    if "anexo 5" in text.lower() or "descrição dos casos especiais" in text.lower():
        collecting = True
        continue

    if not collecting:
        continue

    # Novo título detectado
    if text.istitle() and len(text.split()) <= 6:
        if current_custom:
            result["casos_custom_detalhados"].append(current_custom)
        current_custom = {
            "nome": text,
            "descricao": "",
            "fluxo": "",
            "criterios_aceite": "",
            "requisitos_dependencias": ""
        }
        section = "descricao"
        continue

    if current_custom:
        lower = text.lower()
        if "fluxo" in lower:
            section = "fluxo"
            continue
        elif "critério" in lower or "critérios de aceite" in lower:
            section = "criterios_aceite"
            continue
        elif "requerimentos" in lower or "dependência" in lower:
            section = "requisitos_dependencias"
            continue

        if section:
            current_custom[section] += text + " "

# Não esquecer de salvar o último
if current_custom:
        result["casos_custom_detalhados"].append(current_custom)

    return result









